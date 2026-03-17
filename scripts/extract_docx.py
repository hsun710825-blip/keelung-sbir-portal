from __future__ import annotations

from pathlib import Path

from docx import Document
import sys


def _collapse_ws(s: str) -> str:
    return " ".join(s.split())


def main() -> None:
    # Force UTF-8 output on Windows consoles (avoid cp950 encode crashes).
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # py>=3.7
    except Exception:
        pass

    docx_path = Path(
        r"C:\恂電腦D槽備份\恂\02-業務\05-115\02-115基隆市SBIR\申請文件\115年度SBIR-計畫書內容.docx"
    )
    if not docx_path.exists():
        raise SystemExit(f"Docx not found: {docx_path}")

    doc = Document(str(docx_path))

    print("=== PARAGRAPHS ===")
    for i, para in enumerate(doc.paragraphs, 1):
        t = _collapse_ws(para.text.strip())
        if t:
            print(f"P{i}: {t}")

    print("\n=== TABLES ===")
    for ti, table in enumerate(doc.tables, 1):
        print(f"\n-- TABLE {ti} ({len(table.rows)} rows) --")
        for ri, row in enumerate(table.rows, 1):
            cells = [_collapse_ws(c.text) for c in row.cells]
            print(f"R{ri}: " + " | ".join(cells))


if __name__ == "__main__":
    main()

